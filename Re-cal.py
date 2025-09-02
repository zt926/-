import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import io
import pandas as pd
from docx import Document
from docx.shared import Inches
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors

# ===============================
# 页面配置
# ===============================
st.set_page_config(page_title="Reynolds 数计算器", page_icon="🌊", layout="centered")

# ===============================
# 标题与公式
# ===============================
st.title("🌊 渗流 Reynolds 数计算器")
st.markdown("根据公式:  \n"
            r"$$Re = \frac{v \rho \sqrt{K}}{17.50 \mu \phi^{3/2}}$$")

# ===============================
# 参数输入
# ===============================
st.sidebar.header("请输入参数")

v = st.sidebar.number_input("渗流速度 v (cm/s)", min_value=0.0, value=0.1, format="%.6f")
K = st.sidebar.number_input("渗透率 K (μm²)", min_value=0.0, value=100.0, format="%.6f")
rho = st.sidebar.number_input("密度 ρ (g/cm³)", min_value=0.0, value=1.0, format="%.6f")
mu = st.sidebar.number_input("粘度 μ (mPa·s)", min_value=0.0001, value=1.0, format="%.6f")  # 避免除0
phi = st.sidebar.slider("孔隙度 φ (小数)", 0.0001, 0.8, 0.25, step=0.0001, format="%.4f")

# ===============================
# Re 计算
# ===============================
if v > 0 and K > 0 and rho > 0 and mu > 0 and phi > 0:
    Re = (v * rho * (K ** 0.5)) / (17.50 * mu * (phi ** 1.5))
    st.success(f"✅ 计算结果: Re = **{Re:.6f}**")
else:
    st.warning("⚠️ 请在左侧输入有效的参数以进行计算。")

# ===============================
# 动态曲线绘制 + 下载功能
# ===============================
st.markdown("### 📈 Re 数变化趋势")

option = st.selectbox("选择绘图参数：", ["渗流速度 v", "孔隙度 φ", "双参数热力图 (v-φ)", "双参数等高线图 (v-φ)"])

fig = None
df = None   # 用于存储数据表

if option == "渗流速度 v":
    v_vals = np.linspace(0.01, 2.0, 100)
    Re_vals = (v_vals * rho * (K ** 0.5)) / (17.50 * mu * (phi ** 1.5))
    df = pd.DataFrame({"v (cm/s)": v_vals, "Re": Re_vals})

    fig, ax = plt.subplots()
    ax.plot(v_vals, Re_vals, label=f"φ = {phi:.2f}")
    ax.set_xlabel("渗流速度 v (cm/s)")
    ax.set_ylabel("Reynolds 数 Re")
    ax.set_title("Re 随 v 变化曲线")
    ax.legend()
    st.pyplot(fig)

elif option == "孔隙度 φ":
    phi_vals = np.linspace(0.01, 0.8, 100)
    Re_vals = (v * rho * (K ** 0.5)) / (17.50 * mu * (phi_vals ** 1.5))
    df = pd.DataFrame({"φ": phi_vals, "Re": Re_vals})

    fig, ax = plt.subplots()
    ax.plot(phi_vals, Re_vals, color="orange", label=f"v = {v:.2f}")
    ax.set_xlabel("孔隙度 φ")
    ax.set_ylabel("Reynolds 数 Re")
    ax.set_title("Re 随 φ 变化曲线")
    ax.legend()
    st.pyplot(fig)

elif option == "双参数热力图 (v-φ)" or option == "双参数等高线图 (v-φ)":
    v_vals = np.linspace(0.01, 2.0, 50)
    phi_vals = np.linspace(0.01, 0.8, 50)
    V, PHI = np.meshgrid(v_vals, phi_vals)
    Re_vals = (V * rho * (K ** 0.5)) / (17.50 * mu * (PHI ** 1.5))
    df = pd.DataFrame({
        "v (cm/s)": V.flatten(),
        "φ": PHI.flatten(),
        "Re": Re_vals.flatten()
    })

    fig, ax = plt.subplots()
    if option == "双参数热力图 (v-φ)":
        c = ax.contourf(V, PHI, Re_vals, levels=20, cmap="viridis")
        fig.colorbar(c, ax=ax, label="Reynolds 数 Re")
        ax.set_title("Re 随 v 和 φ 的分布热力图")
    else:
        c = ax.contour(V, PHI, Re_vals, levels=15, cmap="plasma")
        ax.clabel(c, inline=True, fontsize=8)
        ax.set_title("Re 随 v 和 φ 的等高线图")
    ax.set_xlabel("渗流速度 v (cm/s)")
    ax.set_ylabel("孔隙度 φ")
    st.pyplot(fig)

# ===============================
# 下载按钮（PNG + Excel）
# ===============================
if fig is not None:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=300, bbox_inches="tight")
    buf.seek(0)
    st.download_button(
        label="📥 下载图像 (PNG)",
        data=buf,
        file_name=f"Reynolds_{option}.png",
        mime="image/png"
    )

if df is not None:
    excel_buf = io.BytesIO()
    df.to_excel(excel_buf, index=False, sheet_name="Re_Data")
    excel_buf.seek(0)
    st.download_button(
        label="📊 下载数据表 (Excel)",
        data=excel_buf,
        file_name=f"Reynolds_{option}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# ===============================
# 报告导出（Word + PDF）
# ===============================
if st.button("📄 导出 Word 报告"):
    doc = Document()
    doc.add_heading("Reynolds 数计算报告", 0)

    doc.add_heading("输入参数", level=1)
    table = doc.add_table(rows=6, cols=2)
    params = [("v (cm/s)", v), ("K (μm²)", K), ("ρ (g/cm³)", rho),
              ("μ (mPa·s)", mu), ("φ", phi), ("Re", Re)]
    for i, (k, val) in enumerate(params):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = f"{val:.6f}"

    doc.add_heading("公式", level=1)
    doc.add_paragraph("Re = vρ√K / (17.50 μ φ^(3/2))")

    doc.add_heading("趋势图", level=1)
    img_stream = io.BytesIO()
    fig.savefig(img_stream, format="png", dpi=300)
    doc.add_picture(io.BytesIO(img_stream.getvalue()), width=Inches(5))

    doc.add_heading("数据表", level=1)
    doc.add_paragraph(df.head().to_string())

    word_buf = io.BytesIO()
    doc.save(word_buf)
    word_buf.seek(0)
    st.download_button("⬇️ 下载 Word 报告", word_buf, file_name="Reynolds_Report.docx")

if st.button("📑 导出 PDF 报告"):
    pdf_buf = io.BytesIO()
    doc = SimpleDocTemplate(pdf_buf, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("Reynolds 数计算报告", styles['Title']))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("输入参数", styles['Heading1']))
    data = [["参数", "数值"]] + [[k, f"{val:.6f}"] for k, val in [
        ("v (cm/s)", v), ("K (μm²)", K), ("ρ (g/cm³)", rho),
        ("μ (mPa·s)", mu), ("φ", phi), ("Re", Re)]]
    table = Table(data, hAlign='LEFT')
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(table)
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("公式", styles['Heading1']))
    elements.append(Paragraph("Re = vρ√K / (17.50 μ φ^(3/2))", styles['Normal']))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("趋势图", styles['Heading1']))
    img_buf = io.BytesIO()
    fig.savefig(img_buf, format="png", dpi=300)
    img_buf.seek(0)
    elements.append(Image(img_buf, width=400, height=300))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("数据表 (前10行)", styles['Heading1']))
    data_table = [df.columns.tolist()] + df.head(10).values.tolist()
    table = Table(data_table, hAlign='LEFT')
    table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER')
    ]))
    elements.append(table)

    doc.build(elements)
    pdf_buf.seek(0)
    st.download_button("⬇️ 下载 PDF 报告", pdf_buf, file_name="Reynolds_Report.pdf", mime="application/pdf")
